import json
from datetime import datetime
from base64 import b64decode
from os import path, sep, rename
from typing import Callable
import yaml

from bot.message import TextMessageFromDB, MessageType
from utils.msg_pb2 import MessageBytesExtra
from utils.project_path import DATA_PATH
from utils.compress_content_praser import parse_compressed_content
from utils.toolkit import xml_to_dict
from utils.room_data_pb2 import ChatRoomData
from databases.global_config_database import MemoryDatabase
from databases.image_recognition_database import ImageRecognitionDatabase

from requests import post
from docx import Document
from docx.shared import Pt
import xmltodict

# 全局联系人列表伪缓存。
# `process_messages`在循环消息列表时，每次循环会调用`get_talker_name`查询联系人的微信名与备注。
# `get_talker_name`会在首次查询时缓存所有的联系人信息，后续查询直接从缓存中获取。
CONTACT_LIST = {}

def get_all_message(db_handle: list, wxid, include_image=True, start_time=None, end_time=None, port=19001, include_message_type: list=None):
    """
    获取指定联系人的所有消息。
    :param db_handle: MicroMsg.db数据库句柄
    :param wxid: 联系人的wxid
    :param include_image: 是否包含图片消息。默认为 False。
    :param start_time: 起始时间，格式为 'YYYY-MM-DD HH:MM:SS'。默认为 None。
    :param end_time: 结束时间，格式为 'YYYY-MM-DD HH:MM:SS'。默认为 None。
    :param port: 端口号
    :return: 消息列表。
    """
    include_message_type = include_message_type or [
        MessageType.TEXT_MESSAGE,
        MessageType.VOICE_MESSAGE,
        MessageType.VIDEO_MESSAGE,
        MessageType.LOCATION_MESSAGE,
        MessageType.EMOJI_MESSAGE,
        MessageType.IMAGE_MESSAGE,
        MessageType.XML_MESSAGE,
        MessageType.NOTICE_MESSAGE,
        MessageType.CARD_MESSAGE
    ]
    message_type = ', '.join(f'"{item}"' for item in include_message_type)

    # 处理时间戳或格式化时间
    time_condition = ""

    # 转换 start_time
    if start_time:
        if isinstance(start_time, int):  # 如果传入的是时间戳，直接使用
            start_timestamp = start_time
        else:  # 否则，假设传入的是格式化的时间字符串
            try:
                start_timestamp = int(datetime.strptime(start_time, '%Y-%m-%d %H:%M:%S').timestamp())
            except ValueError:
                raise ValueError(f"Invalid start_time format: {start_time}. Expected 'YYYY-MM-DD HH:MM:SS'.")

        time_condition += f" AND CreateTime >= {start_timestamp}"

    # 转换 end_time
    if end_time:
        if isinstance(end_time, int):  # 如果传入的是时间戳，直接使用
            end_timestamp = end_time
        else:  # 否则，假设传入的是格式化的时间字符串
            try:
                end_timestamp = int(datetime.strptime(end_time, '%Y-%m-%d %H:%M:%S').timestamp())
            except ValueError:
                raise ValueError(f"Invalid end_time format: {end_time}. Expected 'YYYY-MM-DD HH:MM:SS'.")

        time_condition += f" AND CreateTime <= {end_timestamp}"

    result = []
    for handle in db_handle:
        # 构建 SQL 查询语句
        body = {
            "dbHandle": handle,
            "sql": f"SELECT * FROM MSG WHERE StrTalker = \"{wxid}\" AND Type in ({message_type}){time_condition} ORDER BY "
                   f"CreateTime ASC;",
        }
        response = (post(f'http://127.0.0.1:{port}/api/execSql', json=body).json()).get('data')
        result += response[1:]

    result.sort(key=lambda x: x[6])
    return result


def xml_message_parse(compressed_content: str):
    result = {
        'content': '[未解析的XML消息]',
        'type': None,
        'ext_info': {}
    }

    # 处理引用回复消息
    def refer_msg(app_msg: dict):
        original_message = app_msg.get('refermsg')
        result['type'] = 'reply'
        # result['content'] = f'回复[{original_message.get("displayname")}]:\n「{original_content}」\n----------\n{app_msg.get("title")}'
        result['content'] = app_msg.get("title")
        result['ext_info']['reply_msg_id'] = original_message.get('svrid')

    def music_share(app_msg: dict):
        result['type'] = 'music_share'
        result['content'] = f'[分享音乐: {app_msg.get("des") or "未知歌手"} - {app_msg.get("title")}]'


    def mini_program(app_msg: dict):
        result['type'] = 'mini_program'
        result['content'] = f'[小程序: {app_msg.get("sourcedisplayname", "未知小程序")}]\n{app_msg.get("title")}'


    def chat_history(app_msg: dict):
        result['type'] = 'chat_history'
        result['content'] = f'[聊天记录：{app_msg.get("title")}]\n{app_msg.get("des")}'

    
    def video_message(app_msg: dict):
        result['type'] = 'video_message'
        result['content'] = f'[视频链接: {app_msg.get("title")}]\n{app_msg.get("des")}'

    def web_view(app_msg: dict):
        result['type'] = 'web_view'
        result['content'] = f'[网页链接: {app_msg.get("title")}]\n{app_msg.get("des")}'
        if app_msg.get('sourcedisplayname'):
            result['content'] += f'\n来源: {app_msg.get("sourcedisplayname")}'

    def default_message(app_msg: dict):
        result['type'] = 'default'
        result['content'] = f'[卡片消息: {app_msg.get("title")}]'

    try:
        parse_result = parse_compressed_content(compressed_content)
        prase_result_dict = xml_to_dict(parse_result)
        app_msg = prase_result_dict.get('msg', {'appmsg': {}})
        app_msg = app_msg.get('appmsg')
        original_msg_type = app_msg.get('type')
        content_dict = {
            "57": refer_msg,
            "3": music_share,
            "92": music_share,
            "33": mini_program,
            "19": chat_history,
            "4": video_message,
            "5": web_view
        } 
        
        content_dict.get(original_msg_type, default_message)(app_msg)

    except Exception as e:
        pass
    return result


def notice_message_parse(content: str):
    if "<revokemsg>" in content:
        content = f'[通知消息: 撤回]\n{content.replace("<revokemsg>", "").replace("</revokemsg>", "")}'
    elif '加入了群聊' in content:
        content = f'[通知消息: 加入群聊]\n{content}'
    elif '拍了拍' in content:
        content = f'[通知消息: 拍一拍]\n{content}'
    else: content = f'[通知消息: 未知类型]\n{content}'
    return content

def check_mention_list(bytes_extra: str):
    """
    检查bytesExtra中是否有@人
    :param bytes_extra: 数据库中bytesExtra
    :return: @人列表, 没有@人返回空列表
    """
    bytes_string = b64decode(bytes_extra)
    msg_bytes_extra = MessageBytesExtra()
    msg_bytes_extra.ParseFromString(bytes_string)
    msg_source = ''
    mention_list = []
    for item in msg_bytes_extra.message2:
        if item.field1 == 7:
            msg_source = item.field2
            break
    if not msg_source:
        return mention_list
    msg_source = f"<container>{msg_source}</container>"
    parse_result = xmltodict.parse(msg_source)
    parse_result = parse_result.get('container')
    msg_source = parse_result.get('msgsource', {})

    at_user_list: str = msg_source.get('atuserlist')
    if not at_user_list:
        return mention_list

    mention_list = at_user_list.split(',')
    return mention_list

def parse_location(content: str):
    try:
        parse_result = xml_to_dict(content)
        location = parse_result.get('msg', {'location': {}}).get('location')
        return f"[位置消息: {location.get('@label')}{location.get('@poiname')}]"
    except Exception as e:
        return '[位置消息]'
    
def card_message_parse(content: str):
    try:
        parse_result = xml_to_dict(content)
        card_info = parse_result.get('msg', {'@nickname': "未知名片"})
        return f"[名片消息: {card_info.get('@nickname')}]"
    except Exception as e:
        print(e)
        return '[名片消息]'

def get_sender_form_room_msg(bytes_extra: str) -> str:
    """
    从bytesExtra中解析消息发送人
    :param bytes_extra: 数据库中bytesExtra
    :return: 发言人的wxid
    """
    bytes_string = b64decode(bytes_extra)
    msg_bytes_extra = MessageBytesExtra()
    msg_bytes_extra.ParseFromString(bytes_string)
    for item in msg_bytes_extra.message2:
        if item.field1 == 1:
            return item.field2
    return ""


def get_memory(from_user, to_user):
    db = MemoryDatabase()
    memories = db.get_memory(
        from_user=from_user,
        to_user=to_user
    )
    return [ { "type": item[1], "content": item[2], "event_time": item[3], "wxid": to_user } for item in memories]
    


def decode_img(message: TextMessageFromDB, save_dir, port=19001, user_data_path: str = None) -> str:
    """
    解码图片
    :param message: 消息对象
    :param save_dir: 保存目录
    :param port: 端口号
    :return: 图片路径
    """
    if message.Type != '3':
        return ""

    user_data_path = user_data_path or post(f'http://127.0.0.1:{port}/api/userInfo').json().get('data').get('dataSavePath')

    message_id = message.MsgSvrID

    post(f'http://127.0.0.1:{port}/api/downloadAttach', json={"msgId": message_id})

    bs64 = message.BytesExtra
    text = b64decode(bs64)

    msg_bytes_extra = MessageBytesExtra()
    msg_bytes_extra.ParseFromString(text)

    image_path = r""

    for item in msg_bytes_extra.message2:
        if item.field1 == 4:
            image_path = item.field2
            break

    dir_path, filename = path.split(image_path)
    image_path = path.join(sep.join(dir_path.split(sep)[1:]), filename)

    body = {
        'filePath': path.join(user_data_path, image_path),
        'storeDir': save_dir
    }
    resp = post(f'http://127.0.0.1:{port}/api/decodeImage', json=body)

    if resp.json().get('code') == 0:
        return ""

    result_file_path = path.join(save_dir, filename.replace('.dat', '.jpg'))
    if path.exists(result_file_path):
        return result_file_path

    dat_result_path = result_file_path.replace('.jpg', '.dat')
    if path.exists(dat_result_path):
        rename(dat_result_path, result_file_path)
        return result_file_path
    return ""


def get_room_members(db_handle: str | int, room_id: str, port=19001):
    body = {
        'sql': f"SELECT RoomData FROM ChatRoom WHERE ChatRoomName=\"{room_id}\"",
        'dbHandle': db_handle
    }
    chat_room_resp = post(f'http://127.0.0.1:{port}/api/execSql', json=body).json()
    chat_room_data = chat_room_resp.get('data')[1] if len(chat_room_resp.get('data')) > 1 else [None]
    chat_room_data = chat_room_data[0]

    chat_room_members = {}

    if chat_room_data:
        chat_room_data = b64decode(chat_room_data)
        chat_room_data_parse = ChatRoomData()
        chat_room_data_parse.ParseFromString(chat_room_data)

        for item in chat_room_data_parse.members:
            room_member_wxid = item.wxID
            room_member_display_name = item.displayName
            chat_room_members[room_member_wxid] = {"display_name": room_member_display_name}

    return chat_room_members


def get_talker_name(db_handle: str | int, wxid, port=19001) -> tuple[str, str, str]:
    """
    从数据库获取微信用户的微信名
    :param db_handle: MicroMsg.db数据库句柄
    :param wxid: 微信id
    :param port: 端口号
    :return: tuple[备注, 微信名, WXID]
    """
    contact_list = CONTACT_LIST.get(port, None)

    if not contact_list:
        body = {
            "dbHandle": db_handle,
            # "sql": f"SELECT ct.Remark, ct.NickName, ct.LabelIDList, ct.PYInitial, ct.QuanPin, ct.Reserved1, ct.Reserved2, ct.VerifyFlag, ct.Type, ct.ExtraBuf, cth.bigHeadImgUrl, cth.smallHeadImgUrl FROM Contact AS ct LEFT JOIN ContactHeadImgUrl AS cth ON ct.UserName = cth.usrName"
            "sql": "SELECT ct.UserName, ct.Remark, ct.NickName FROM Contact AS ct LEFT JOIN ContactHeadImgUrl AS cth ON ct.UserName = cth.usrName;"
        }
        resp = post(f'http://127.0.0.1:{port}/api/execSql', json=body).json()

        if resp.get('code') != 1:
            return '', '', ""
                    
        contacts = resp.get('data')
        contacts_dict = {}
        for index, item in enumerate(contacts):
            if index == 0:
                continue
            contacts_dict[item[0]] = item

        CONTACT_LIST[port] = contacts_dict
        contact_list = contacts_dict

    result = contact_list.get(wxid, [])
    remark, nick_name = result[1] if len(result) > 1 else "未知用户", result[2] if len(result) > 2 else "未知用户"
    return remark, nick_name, wxid


def process_messages(msg_db_handle: list, micro_msg_db_handle: str | int, wxid, write_function: Callable,
                     include_image=False, start_time=None, end_time=None,
                     port=19001):
    """
    处理消息
    :param msg_db_handle: msg.db数据库句柄
    :param micro_msg_db_handle: MicroMsg.db数据库句柄
    :param wxid: 联系人的wxid
    :param write_function: 具体写入的回调函数
    :param include_image: 是否包含图片
    :param start_time:
    :param end_time:
    :param port:
    :return:
    """
    data = get_all_message(msg_db_handle, wxid, include_image, port=port, start_time=start_time, end_time=end_time)

    user_info = post(f'http://127.0.0.1:{port}/api/userInfo').json().get('data')

    if not path.exists(DATA_PATH):
        from os import makedirs
        makedirs(DATA_PATH)

    exports_path = path.join(DATA_PATH, 'exports')
    if not path.exists(exports_path):
        from os import makedirs
        makedirs(exports_path)
    is_room = '@chatroom' in wxid

    room_members = {}
    if is_room:
        room_members = get_room_members(db_handle=micro_msg_db_handle, room_id=wxid, port=port)
    
    for index, item in enumerate(data):
        message = TextMessageFromDB(*item)
        # 获取发送人名称
        image_path = ""

        if message.IsSender == '1':
            remark, nick_name, sender_id = "用户自己", user_info.get('name'), user_info.get('wxid')
        else:
            sender_id = get_sender_form_room_msg(message.BytesExtra) if message.room else message.StrTalker
            remark, nick_name, _ = get_talker_name(micro_msg_db_handle, sender_id, port)

        if is_room:
            member_info = room_members.get(sender_id, {'display_name': ''})
            member_remark = member_info.get('display_name', remark)
            remark = member_remark if member_remark else remark

        room = message.room
        mention_list = ""

        if room and '@' in message.StrContent:
            # 获取提及人名称
            mention_list = check_mention_list(message.BytesExtra)
            mention_list = [get_talker_name(micro_msg_db_handle, user_id, port) for user_id in mention_list if user_id]
            mention_list = [{"name": _nick_name, "wxid": _wxid} for _remark, _nick_name, _wxid in
                            mention_list]

        if message.Type == MessageType.IMAGE_MESSAGE and include_image:
            image_path = decode_img(message, path.join(DATA_PATH, 'images'), port=port)

        format_time = datetime.fromtimestamp(int(message.CreateTime)).strftime('%Y-%m-%d %H:%M:%S')
        message_content = image_path if message.Type == MessageType.IMAGE_MESSAGE and include_image else message.StrContent
        write_function(nick_name, remark, format_time, message_content, mention_list, room, message, sender_id)


def write_doc(msg_db_handle: list, micro_msg_db_handle: str | int, wxid, doc_filename=None, include_image=False,
              port=19001, start_time=None, end_time=None):
    """
    将聊天记录写入docx文件中
    :param msg_db_handle:
    :param micro_msg_db_handle:
    :param wxid:
    :param doc_filename:
    :param include_image:
    :param port:
    :param start_time:
    :param end_time:
    :return:
    """
    doc = Document()

    main_remark, main_username, _ = get_talker_name(micro_msg_db_handle, wxid, port=port)
    doc_filename = f"{main_username}_{main_remark}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx" if not \
        doc_filename else doc_filename

    is_room = '@chatroom' in wxid

    doc.add_paragraph()
    doc.add_heading("数据说明：", level=1)
    doc.add_paragraph(f'这是一份微信{"群聊" if is_room else "私聊"}聊天记录，其中：')
    doc.add_paragraph('每条消息开始标记为 `===消息开始===`', style='ListBullet')
    doc.add_paragraph('每条消息结束标记为 `===消息结束===`', style='ListBullet')
    doc.add_paragraph('发送人代表消息发送人，其中同名的发送人代表同一个人发送的消息', style='ListBullet')
    doc.add_paragraph('备注代表消息发送人的备注信息，如果没有备注字段则代表对该联系人没有备注。', style='ListBullet')
    if is_room: doc.add_paragraph('提及人代表该消息提及到的人，为空则没有提及到任何人。', style='ListBullet')
    doc.add_paragraph('内容代表该条消息的内容', style='ListBullet')
    doc.add_paragraph('时间代表消息发送的时间', style='ListBullet')

    doc.add_paragraph()
    doc.add_heading("数据：", level=1)

    def callback(_nick_name, _remark, _format_time, _message_content, _mention_list, _room,
                 _original_message: TextMessageFromDB, sender_id=None):
        doc.add_paragraph("=== 消息开始 ===")
        doc.add_paragraph(f'发送人：{_nick_name}')
        if _remark: doc.add_paragraph(f'备注：{_remark}')
        if _room: doc.add_paragraph(f'提及人：{_mention_list}')
        doc.add_paragraph(f'时间：{_format_time}')

        if _original_message.Type == MessageType.TEXT_MESSAGE:
            doc.add_paragraph(f'内容：\n{_message_content}')
            doc.add_paragraph("=== 消息结束 ===")
            return

        if _original_message.Type == MessageType.IMAGE_MESSAGE and path.exists(_message_content):
            doc.add_paragraph(f'内容：')
            try:
                doc.add_picture(_message_content, width=Pt(300))
            except Exception as e:
                print(e)
                doc.add_paragraph(f'图片添加失败')
        else:
            doc.add_paragraph(f'内容：\n图片获取失败')

        doc.add_paragraph("=== 消息结束 ===")

    process_messages(
        msg_db_handle=msg_db_handle,
        micro_msg_db_handle=micro_msg_db_handle,
        wxid=wxid,
        write_function=callback,
        port=port,
        include_image=include_image, start_time=start_time, end_time=end_time
    )

    file_name = path.join(DATA_PATH, 'exports', doc_filename)
    doc.save(file_name)

    return file_name



def write_txt(msg_db_handle: list, micro_msg_db_handle: str | int, wxid, filename=None,
              port=19001, file_type='json', endswith_txt=True, start_time=None, end_time=None, include_image=False):
    
    user_info: dict = post(f'http://127.0.0.1:{port}/api/userInfo').json().get('data')
    main_remark, main_username, _ = get_talker_name(micro_msg_db_handle, wxid, port=port)
    is_room = '@chatroom' in wxid
    memories = get_memory(from_user=user_info.get('wxid'), to_user=wxid)
    result = {
        "meta": {
            "description": f'聊天记录的数据结构定义',
            "notes": f'这是一份微信{"群聊" if is_room else "私聊"}聊天记录',
            "field": {
                "sender": "消息发送人的名称。，其中同名的发送人代表同一个人发送的消息",
                "remark": "对消息发送人的备注，如果为空则代表该联系人没有备注",
                "content": """
消息的具体内容, 主要包含以下形式：
1. 文本消息：直接展示内容。
    例如："你好"
2. 特殊消息：由`[主类型: 子类型||标题]\n额外描述`组成。
    例如：
        - "[网页链接: 香港最新真实收入曝光！]\n来源: 景鸿移民"。
        - "[小程序: 瑞幸咖啡]\n来杯咖啡..."
        - "[分享音乐: 陶喆 - 爱我还是他]"
        - "[聊天记录: 群聊的聊天记录]\n张三: 你好\n李四: 你也好"
        - "[视频链接: 林俊杰《起风了》]\nUP主：大虾试车真香\n播放：50.3万"
        - "[位置消息: 深圳市南山区xxxxx]"
        - "[名片消息: 张三]"
        - "[引用消息：张三 回复 李四]\n原始消息(部分): 「今天天气真好好！」\n回复内容: 是啊！"
        - "[通知消息：拍一拍]\n张三 拍了拍 李四" || "[通知消息: 撤回]\n张三撤回了一条消息" || "[通知消息: 邀请]\n张三邀请李四进入群聊"
""",
                "time": "消息发送时间",
                "wxid": "消息发送人的wxid，每个用户的唯一id，可以用来判断消息发送人是否是同一位。",
                "msg_id": "消息的唯一ID",
                "reply_msg_id": "当这条消息引用(回复)了另一条消息，则存放被引用(回复)的消息的msg_id，否则不展示这个字段。"
            },
            "context": {
                "description": "AI生成辅助上下文（根据历史消息动态推断得出），仅用于理解对话隐含信息，禁止直接输出到最终结论中。",
                "memories": memories
            }
        },
        "data": []
    }

    if is_room: result['meta']['field']['mentioned'] = "消息中提及到的用户，如果为空则代表这条消息没有提及任何人。格式为：[{'name': '被提及人名称', 'wxid': '被提及人wxid'}]"
    
    image_rec_db = ImageRecognitionDatabase()

    def callback(_nick_name, _remark, _format_time, _message_content, _mention_list, _room,
                 _original_message: TextMessageFromDB, sender_id=None):
        
        # 根据消息类型，获取对应的内容。表情包和图片描述解析违禁风险过大，暂时不做。
        content_types = {
            MessageType.IMAGE_MESSAGE: "[图片]\n图片描述: 无具体描述",    # 可以通过GML 4V Flash描述图片，然后单独开一个本地db，把描述和msg_id关联。考虑到成本问题，暂时不做。 
            MessageType.VIDEO_MESSAGE: "[视频]",
            MessageType.TEXT_MESSAGE: _message_content,
            MessageType.VOICE_MESSAGE: "[语音]",
            MessageType.LOCATION_MESSAGE: "[位置消息]",
            MessageType.EMOJI_MESSAGE: "[动画表情]",    # 同样可以通过content字段的cndurl下载表情图片，GML 4V Flash描述图片用本地db存起来用标签的md5作为id，同样考虑成本问题暂时不做。
            MessageType.XML_MESSAGE: "[未解析的XML消息]",
            MessageType.NOTICE_MESSAGE: "[通知消息]",
        }

        reply_msg_id = None

        if _original_message.Type == MessageType.XML_MESSAGE:
            xml_prese_result = xml_message_parse(_original_message.CompressContent)
            content_types[MessageType.XML_MESSAGE] = xml_prese_result.get('content', '[未解析的XML消息]')
            reply_msg_id = xml_prese_result.get('ext_info', {'reply_msg_id', None}).get('reply_msg_id')
            if reply_msg_id:
                for history in result['data'][::-1]:
                    if history['msg_id'] == reply_msg_id:
                        original_content = history.get("content") if len(history.get("content")) < 10 else f'{history.get("content")[0:5]} ...'
                        content_types[MessageType.XML_MESSAGE] = f"[引用消息：{_nick_name} 回复 {history.get('sender')}]\n原始消息(部分): 「{original_content}」\n回复内容(完整): {content_types[MessageType.XML_MESSAGE]}"
                        break
        
        elif _original_message.Type == MessageType.NOTICE_MESSAGE:
            content_types[MessageType.NOTICE_MESSAGE] = notice_message_parse(_original_message.content)
            # 通知消息会有一部分消息获取不到wxid和名称，用微信团队兜底
            if _nick_name == "未知用户" and not sender_id:
                _nick_name, _remark = "微信团队", ""
                sender_id = "weixin"

        elif _original_message.Type == MessageType.CARD_MESSAGE:
            content_types[MessageType.CARD_MESSAGE] = card_message_parse(_original_message.content)

        elif _original_message.Type == MessageType.LOCATION_MESSAGE:
            content_types[MessageType.LOCATION_MESSAGE] = parse_location(_message_content)

        elif _original_message.Type == MessageType.IMAGE_MESSAGE:
            _, recognition_result, _ = image_rec_db.get_recognition_result(_original_message.MsgSvrID)
            if recognition_result is not None:
                content_types[MessageType.IMAGE_MESSAGE] = f"[图片]\n图片描述: {recognition_result}"

        item = {
            "sender": _nick_name,
            "remark": _remark,
            "content": content_types.get(_original_message.Type),
            "time": _format_time,
            "wxid": sender_id,
            "msg_id": _original_message.MsgSvrID,
        }
        
        if _room and _mention_list: item['mentioned'] = _mention_list
        if reply_msg_id: item['reply_msg_id'] = reply_msg_id

        result['data'].append(item)


    process_messages(
        msg_db_handle=msg_db_handle,
        micro_msg_db_handle=micro_msg_db_handle,
        wxid=wxid,
        write_function=callback,
        port=port,
        include_image=include_image, start_time=start_time, end_time=end_time
    )

    if file_type is None: return result

    if file_type.lower() not in ['json', 'yaml', 'yml']: file_type = 'json'

    file_ends = '.txt' if endswith_txt else f".{file_type}"
    filename = f"{main_username}_{main_remark}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}{file_ends}" if not \
        filename else filename
    file_path = path.join(DATA_PATH, 'exports', filename)
    if file_type.lower() == 'json':
        with open(file_path, 'w', encoding='utf-8') as fw:
            fw.write(json.dumps(result, ensure_ascii=False, indent=4))
        return file_path

    if file_type.lower() == 'yaml' or file_type.lower() == 'yml':
        with open(file_path, 'w', encoding='utf-8') as fw:
            fw.write(yaml.dump({"meta": result.get('meta')}, allow_unicode=True))

        with open(file_path, 'a', encoding='utf-8') as fa:
            fa.write('\n')
            fa.write(yaml.dump({"data": result.get('data')}, allow_unicode=True))
        return file_path
    

