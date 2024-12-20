import json
import re
from datetime import datetime
from base64 import b64decode
from os import path, sep

from message import TextMessageFromDB

from requests import post
from docx import Document
from docx.shared import Pt

DATA_PATH = post('http://127.0.0.1:19001/api/userInfo').json().get('data').get('dataSavePath')


def get_all_message(db_handle, wxid, include_image):
    message_type = "(\"3\", \"1\")" if include_image else "(\"1\")"
    body = {
        "dbHandle": db_handle,
        "sql": f"SELECT * FROM MSG WHERE StrTalker = \"{wxid}\" AND Type in {message_type} ORDER BY CreateTime DESC",
    }
    return (post('http://127.0.0.1:19001/api/execSql', json=body).json()).get('data')


def decode_img(message: TextMessageFromDB, save_dir) -> str:
    if message.Type != '3':
        return
    message_id = message.MsgSvrID

    post('http://127.0.0.1:19001/api/downloadAttach', json={"msgId": message_id})

    bs64 = message.BytesExtra
    text = b64decode(bs64).decode('mac_roman')
    pattern = r'([a-zA-Z0-9_\\/:.-]+(?:\\|/)[a-zA-Z0-9_\\/:.-]*\.dat)'
    img_paths = re.findall(pattern, text)
    image_path: str = img_paths[0]
    dir_path, filename = path.split(image_path)
    image_path = path.join(sep.join(dir_path.split(sep)[1:]), filename)

    body = {
        'filePath': path.join(DATA_PATH, image_path),
        'storeDir': save_dir
    }
    resp = post('http://127.0.0.1:19001/api/decodeImage', json=body)
    if resp.json().get('code') == 0:
        return ""
    return path.join(save_dir, filename.replace('.dat', '.jpg'))


def write_doc(db_handle, wxid, doc_filename='chat_record.docx', include_image=False):
    data = get_all_message(db_handle, wxid, include_image)
    doc = Document()

    for index, item in enumerate(data):
        if index == 0:
            continue

        message = TextMessageFromDB(*item)
        image_path = ""

        if message.Type == '3':
            image_path = decode_img(message, r'D:\wangyingjie\WeBot\config')

        temp = {
            "role": "other" if message.IsSender == "0" else "me",
            "content": message.StrContent if message.Type == '1' else image_path,
            "time": datetime.fromtimestamp(int(message.CreateTime)).strftime('%Y-%m-%d %H:%M:%S'),
            "type": "text" if message.Type == "1" else "image"
        }

        if temp['type'] == "image" and temp['content'] == '':
            continue

        doc.add_paragraph(f'角色：{"我" if temp["role"] == "me" else "对方"}')
        doc.add_paragraph(f'时间：{temp["time"]}')

        if temp["type"] == "text":
            doc.add_paragraph(f'内容：\n{temp["content"]}')

        if temp["type"] == "image":
            doc.add_paragraph(f'内容：')
            doc.add_picture(image_path, width=Pt(300))
        doc.add_paragraph('---')

    doc.save(doc_filename)


if __name__ == '__main__':
    write_doc(
        1535143968448,
        'wxid_bymczps415pj22'
    )
