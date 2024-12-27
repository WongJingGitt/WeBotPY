import json
from datetime import datetime
from base64 import b64decode
from os import path, sep, rename
from typing import Callable
import yaml

from libs.message import TextMessageFromDB, MessageType
from utils.msg_pb2 import MessageBytesExtra
from utils.toolkit import DATA_PATH

from requests import post
from docx import Document
from docx.shared import Pt
import xmltodict


def get_all_message(db_handle, wxid, include_image, start_time=None, end_time=None, port=19001):
    image_type = MessageType.IMAGE_MESSAGE if include_image else ""
    message_type = f"\"{MessageType.TEXT_MESSAGE}\", \"{MessageType.VOICE_MESSAGE}\", \"{MessageType.VIDEO_MESSAGE}\", \"{MessageType.LOCATION_MESSAGE}\", \"{MessageType.EMOJI_MESSAGE}\""

    if image_type:
        message_type += f", \"{image_type}\""

    # 处理时间戳或格式化时间
    time_condition = ""

    # 转换 start_time
    if start_time:
        if isinstance(start_time, int):  # 如果传入的是时间戳，直接使用
            start_timestamp = start_time
        else:  # 否则，假设传入的是格式化的时间字符串
            try:
                start_timestamp = int(datetime.strptime(start_time, '%Y-%m-%d %H:%M:%S').timestamp())
            except ValueError:
                raise ValueError(f"Invalid start_time format: {start_time}. Expected 'YYYY-MM-DD HH:MM:SS'.")

        time_condition += f" AND CreateTime >= {start_timestamp}"

    # 转换 end_time
    if end_time:
        if isinstance(end_time, int):  # 如果传入的是时间戳，直接使用
            end_timestamp = end_time
        else:  # 否则，假设传入的是格式化的时间字符串
            try:
                end_timestamp = int(datetime.strptime(end_time, '%Y-%m-%d %H:%M:%S').timestamp())
            except ValueError:
                raise ValueError(f"Invalid end_time format: {end_time}. Expected 'YYYY-MM-DD HH:MM:SS'.")

        time_condition += f" AND CreateTime <= {end_timestamp}"

    # 构建 SQL 查询语句
    body = {
        "dbHandle": db_handle,
        "sql": f"SELECT * FROM MSG WHERE StrTalker = \"{wxid}\" AND Type in ({message_type}){time_condition} ORDER BY "
               f"CreateTime ASC",
    }

    return (post(f'http://127.0.0.1:{port}/api/execSql', json=body).json()).get('data')


def check_mention_list(bytes_extra: str):
    bytes_string = b64decode(bytes_extra)
    msg_bytes_extra = MessageBytesExtra()
    msg_bytes_extra.ParseFromString(bytes_string)
    msg_source = ''
    mention_list = []
    for item in msg_bytes_extra.message2:
        if item.field1 == 7:
            msg_source = item.field2
            break
    if not msg_source:
        return mention_list
    msg_source = f"<container>{msg_source}</container>"
    parse_result = xmltodict.parse(msg_source)
    parse_result = parse_result.get('container')
    msg_source = parse_result.get('msgsource', {})

    at_user_list: str = msg_source.get('atuserlist')
    if not at_user_list:
        return mention_list

    mention_list = at_user_list.split(',')
    return mention_list


def get_sender_form_room_msg(bytes_extra: str) -> str:
    """
    从bytesExtra中解析消息发送人
    :param bytes_extra: 数据库中bytesExtra
    :return: 发言人的wxid
    """
    bytes_string = b64decode(bytes_extra)
    msg_bytes_extra = MessageBytesExtra()
    msg_bytes_extra.ParseFromString(bytes_string)
    for item in msg_bytes_extra.message2:
        if item.field1 == 1:
            return item.field2
    return ""


def decode_img(message: TextMessageFromDB, save_dir, port=19001) -> str:
    if message.Type != '3':
        return ""

    user_data_path = post(f'http://127.0.0.1:{port}/api/userInfo').json().get('data').get('dataSavePath')

    message_id = message.MsgSvrID

    post(f'http://127.0.0.1:{port}/api/downloadAttach', json={"msgId": message_id})

    bs64 = message.BytesExtra
    text = b64decode(bs64)

    msg_bytes_extra = MessageBytesExtra()
    msg_bytes_extra.ParseFromString(text)

    image_path = r""

    for item in msg_bytes_extra.message2:
        if item.field1 == 4:
            image_path = item.field2
            break

    dir_path, filename = path.split(image_path)
    image_path = path.join(sep.join(dir_path.split(sep)[1:]), filename)

    body = {
        'filePath': path.join(user_data_path, image_path),
        'storeDir': save_dir
    }
    resp = post(f'http://127.0.0.1:{port}/api/decodeImage', json=body)

    if resp.json().get('code') == 0:
        return ""

    result_file_path = path.join(save_dir, filename.replace('.dat', '.jpg'))
    if path.exists(result_file_path):
        return result_file_path

    dat_result_path = result_file_path.replace('.jpg', '.dat')
    if path.exists(dat_result_path):
        rename(dat_result_path, result_file_path)
        return result_file_path
    return ""


def get_talker_name(db_handle, wxid, port=19001) -> tuple[str, str]:
    """
    从数据库获取微信用户的微信名
    :param db_handle: MicroMsg.db数据库句柄
    :param wxid: 微信id
    :param port: 端口号
    :return: tuple[备注, 微信名]
    """
    body = {
        "dbHandle": db_handle,
        "sql": f"SELECT Remark,NickName From Contact Where UserName = \"{wxid}\""
    }
    resp = post(f'http://127.0.0.1:{port}/api/execSql', json=body).json()
    if resp.get('code') != 1:
        return '', ''

    result = resp.get('data')
    if len(result) < 2:
        return '', ''

    remark, nick_name = result[-1]
    return remark, nick_name


def process_messages(msg_db_handle, micro_msg_db_handle, wxid, write_function: Callable, filename=None,
                     include_image=False, start_time=None, end_time=None,
                     port=19001):
    data = get_all_message(msg_db_handle, wxid, include_image, port=port, start_time=start_time, end_time=end_time)

    user_info = post(f'http://127.0.0.1:{port}/api/userInfo').json().get('data')

    if not path.exists(DATA_PATH):
        from os import makedirs
        makedirs(DATA_PATH)

    exports_path = path.join(DATA_PATH, 'exports')
    if not path.exists(exports_path):
        from os import makedirs
        makedirs(exports_path)

    for index, item in enumerate(data):
        if index == 0:
            continue

        message = TextMessageFromDB(*item)
        # 获取发送人名称
        image_path = ""

        if message.IsSender == '1':
            remark, nick_name = "用户自己", user_info.get('name')
        else:
            sender_id = get_sender_form_room_msg(message.BytesExtra) if message.room else message.StrTalker
            remark, nick_name = get_talker_name(micro_msg_db_handle, sender_id, port)

        room = message.room
        mention_list = ""

        if room:
            # 获取提及人名称
            mention_list = check_mention_list(message.BytesExtra)
            mention_list = [get_talker_name(micro_msg_db_handle, user_id, port) for user_id in mention_list if user_id]
            mention_list = [f"{_nick_name}({_remark})" if _remark else _nick_name for _remark, _nick_name in
                            mention_list]

        if message.Type == MessageType.IMAGE_MESSAGE:
            image_path = decode_img(message, path.join(DATA_PATH, 'images'), port=port)

        mention_list = ' || '.join(mention_list)
        format_time = datetime.fromtimestamp(int(message.CreateTime)).strftime('%Y-%m-%d %H:%M:%S')
        message_content = message.StrContent if message.Type == MessageType.TEXT_MESSAGE else image_path

        write_function(nick_name, remark, format_time, message_content, mention_list, room, message)


def write_doc(msg_db_handle, micro_msg_db_handle, wxid, doc_filename=None, include_image=False,
              port=19001, start_time=None, end_time=None):
    """
    将聊天记录写入docx文件中
    :param msg_db_handle:
    :param micro_msg_db_handle:
    :param wxid:
    :param doc_filename:
    :param include_image:
    :param port:
    :param start_time:
    :param end_time:
    :return:
    """
    doc = Document()

    main_remark, main_username = get_talker_name(micro_msg_db_handle, wxid, port=port)
    doc_filename = f"{main_username}_{main_remark}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx" if not \
        doc_filename else doc_filename

    is_room = '@chatroom' in wxid

    doc.add_paragraph()
    doc.add_heading("数据说明：", level=1)
    doc.add_paragraph(f'这是一份微信{"群聊" if is_room else "私聊"}聊天记录，其中：')
    doc.add_paragraph('每条消息开始标记为 `===消息开始===`', style='ListBullet')
    doc.add_paragraph('每条消息结束标记为 `===消息结束===`', style='ListBullet')
    doc.add_paragraph('发送人代表消息发送人，其中同名的发送人代表同一个人发送的消息', style='ListBullet')
    doc.add_paragraph('备注代表消息发送人的备注信息，如果没有备注字段则代表对该联系人没有备注。', style='ListBullet')
    if is_room: doc.add_paragraph('提及人代表该消息提及到的人，为空则没有提及到任何人。', style='ListBullet')
    doc.add_paragraph('内容代表该条消息的内容', style='ListBullet')
    doc.add_paragraph('时间代表消息发送的时间', style='ListBullet')

    doc.add_paragraph()
    doc.add_heading("数据：", level=1)

    def callback(_nick_name, _remark, _format_time, _message_content, _mention_list, _room,
                 _original_message: TextMessageFromDB):
        doc.add_paragraph("=== 消息开始 ===")
        doc.add_paragraph(f'发送人：{_nick_name}')
        if _remark: doc.add_paragraph(f'备注：{_remark}')
        if _room: doc.add_paragraph(f'提及人：{_mention_list}')
        doc.add_paragraph(f'时间：{_format_time}')

        if _original_message.Type == MessageType.TEXT_MESSAGE:
            doc.add_paragraph(f'内容：\n{_message_content}')
            doc.add_paragraph("=== 消息结束 ===")
            return

        if _original_message.Type == MessageType.IMAGE_MESSAGE and path.exists(_message_content):
            doc.add_paragraph(f'内容：')
            try:
                doc.add_picture(_message_content, width=Pt(300))
            except Exception as e:
                print(e)
                doc.add_paragraph(f'图片添加失败')
        else:
            doc.add_paragraph(f'内容：\n图片获取失败')

        doc.add_paragraph("=== 消息结束 ===")

    process_messages(
        msg_db_handle=msg_db_handle,
        micro_msg_db_handle=micro_msg_db_handle,
        wxid=wxid,
        write_function=callback,
        port=port,
        include_image=include_image, start_time=start_time, end_time=end_time
    )

    file_name = path.join(DATA_PATH, 'exports', doc_filename)
    doc.save(file_name)

    return file_name


def write_txt(msg_db_handle, micro_msg_db_handle, wxid, filename=None,
              port=19001, file_type='json', endswith_txt=True, start_time=None, end_time=None):
    main_remark, main_username = get_talker_name(micro_msg_db_handle, wxid, port=port)
    is_room = '@chatroom' in wxid

    result = {
        "meta": {
            "description": f'聊天记录的数据结构定义',
            "notes": f'这是一份微信{"群聊" if is_room else "私聊"}聊天记录',
            "field": {
                "sender": "消息发送人，其中同名的发送人代表同一个人发送的消息",
                "remark": "对消息发送人的备注，如果为空则代表该联系人没有备注",
                "content": "具体的消息内容",
                "time": "消息发送时间",
            }
        },
        "data": []
    }

    if is_room: result['meta']['field']['mentioned'] = "消息中提及到的用户，如果为空则代表这条消息没有提及任何人"

    def callback(_nick_name, _remark, _format_time, _message_content, _mention_list, _room,
                 _original_message: TextMessageFromDB):

        content_types = {
            MessageType.IMAGE_MESSAGE: "[图片]",
            MessageType.VIDEO_MESSAGE: "[视频]",
            MessageType.TEXT_MESSAGE: _message_content,
            MessageType.VOICE_MESSAGE: "[语音]",
            MessageType.LOCATION_MESSAGE: "[位置]",
            MessageType.EMOJI_MESSAGE: "[动画表情]",
        }

        item = {
            "sender": _nick_name,
            "remark": _remark,
            "content": content_types.get(_original_message.Type),
            "time": _format_time,
        }
        if _room: item['mentioned'] = _mention_list
        result['data'].append(item)

    process_messages(
        msg_db_handle=msg_db_handle,
        micro_msg_db_handle=micro_msg_db_handle,
        wxid=wxid,
        write_function=callback,
        port=port,
        include_image=True, start_time=start_time, end_time=end_time
    )

    if file_type.lower() not in ['json', 'yaml', 'yml']: file_type = 'json'

    file_ends = '.txt' if endswith_txt else f".{file_type}"
    filename = f"{main_username}_{main_remark}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}{file_ends}" if not \
        filename else filename
    file_path = path.join(DATA_PATH, 'exports', filename)

    if file_type.lower() == 'json':
        with open(file_path, 'w', encoding='utf-8') as fw:
            fw.write(json.dumps(result, ensure_ascii=False, indent=4))
        return file_path

    if file_type.lower() == 'yaml' or file_type.lower() == 'yml':
        with open(file_path, 'w', encoding='utf-8') as fw:
            fw.write(yaml.dump({"meta": result.get('meta')}, allow_unicode=True))

        with open(file_path, 'a', encoding='utf-8') as fa:
            fa.write('\n')
            fa.write(yaml.dump({"data": result.get('data')}, allow_unicode=True))
        return file_path
